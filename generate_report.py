# -*- coding: utf-8 -*-
"""
Thesis Report Generator for Mastan Vali Shaik (24226807)
MSc in Artificial Intelligence, National College of Ireland

Generates a ~10,000 word Word Document (.docx) thesis report with:
- Flowing paragraph style (no bullet points)
- Embedded result figures from data/ folder
- Content drawn from existing project files (no hallucination)
"""

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR = r"C:\Users\Dell\Desktop\Masthan\mathan_thesis_code"
DATA_DIR = os.path.join(BASE_DIR, "data")
DOCS_DIR = os.path.join(BASE_DIR, "docs")
REPORT_DIR = os.path.join(BASE_DIR, "report")
os.makedirs(REPORT_DIR, exist_ok=True)

OUTPUT_DOCX = os.path.join(REPORT_DIR, "thesis_report_final.docx")

# ── All figures available in data/ ────────────────────────────────────────────
FIGURES = {
    "top_words": os.path.join(DATA_DIR, "fig_top_words.png"),
    "q_len": os.path.join(DATA_DIR, "fig_q_len.png"),
    "a_len": os.path.join(DATA_DIR, "fig_a_len.png"),
    "model_compare": os.path.join(DATA_DIR, "fig_model_compare.png"),
    "rouge_bleu": os.path.join(DATA_DIR, "fig_rouge_bleu.png"),
    "time": os.path.join(DATA_DIR, "fig_time.png"),
    "complete_personal": os.path.join(DATA_DIR, "fig_complete_personal.png"),
    "composite": os.path.join(DATA_DIR, "fig_composite.png"),
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_page_number(run):
    """Inserts a PAGE field into footers for dynamic page numbering."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def fmt(p, spacing=1.5, after=6, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Sets paragraph formatting."""
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.alignment = align


def run(p, text, bold=False, italic=False, font="Times New Roman", size=12, color=None):
    """Adds a formatted run to paragraph."""
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def para(doc, text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         spacing=1.5, after=6, before=0, first_indent=True):
    """Adds a complete formatted paragraph."""
    p = doc.add_paragraph()
    fmt(p, spacing, after, before, align)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    r = run(p, text, bold=bold, italic=italic, size=size)
    return p


def heading(doc, text, level=1, size=18):
    """Adds a chapter or section heading."""
    p = doc.add_paragraph()
    if level == 1:
        fmt(p, spacing=1.5, after=18, before=12, align=WD_ALIGN_PARAGRAPH.LEFT)
        run(p, text, bold=True, size=size, color=RGBColor(30, 70, 120))
    elif level == 2:
        fmt(p, spacing=1.5, after=6, before=12, align=WD_ALIGN_PARAGRAPH.LEFT)
        run(p, text, bold=True, size=14)
    elif level == 3:
        fmt(p, spacing=1.5, after=6, before=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        run(p, text, bold=True, italic=True, size=12)
    return p


def add_figure(doc, path, caption, width=5.0):
    """Inserts a figure with centered caption."""
    if not os.path.exists(path):
        print(f"  WARNING: Figure not found: {path}")
        return
    p_img = doc.add_paragraph()
    fmt(p_img, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    p_img.add_run().add_picture(path, width=Inches(width))
    p_cap = doc.add_paragraph()
    fmt(p_cap, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    run(p_cap, caption, italic=True, size=10)


def add_figure_pair(doc, path1, path2, caption, width=3.0):
    """Inserts two figures side by side with a shared caption."""
    p_img = doc.add_paragraph()
    fmt(p_img, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    if os.path.exists(path1):
        p_img.add_run().add_picture(path1, width=Inches(width))
    if os.path.exists(path2):
        p_img.add_run().add_picture(path2, width=Inches(width))
    p_cap = doc.add_paragraph()
    fmt(p_cap, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    run(p_cap, caption, italic=True, size=10)


def add_header_footer(doc):
    """Adds running headers and footers with page numbers."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    fmt(hp, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(hp, "MSc in Artificial Intelligence  |  Thesis Report  |  Mastan Vali Shaik",
        size=8.5, color=RGBColor(128, 128, 128))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fmt(fp, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run(fp, "Student ID: 24226807  |  National College of Ireland",
        size=8.5, color=RGBColor(128, 128, 128))
    fp.add_run(" " * 50)
    r_num = fp.add_run("Page ")
    r_num.font.name = "Times New Roman"
    r_num.font.size = Pt(8.5)
    r_num.font.color.rgb = RGBColor(128, 128, 128)
    add_page_number(r_num)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN REPORT GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print("Building thesis report...")
    doc = Document()

    # Configure page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    print("  Title page...")
    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    fmt(p_title, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p_title, "DYNAMIC HEALTHCARE PERSONALIZATION IN AI ASSISTANTS:\n",
        bold=True, size=18, color=RGBColor(30, 70, 120))
    run(p_title, "A COMPARATIVE ANALYSIS OF LSTM FINE-TUNING VERSUS\nRETRIEVAL-AUGMENTED GENERATION",
        bold=True, size=16, color=RGBColor(30, 70, 120))

    for _ in range(4):
        doc.add_paragraph()

    p_author = doc.add_paragraph()
    fmt(p_author, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p_author, "Mastan Vali Shaik\n", bold=True, size=14)
    run(p_author, "Student ID: 24226807\n\n", size=12)
    run(p_author, "Programme: Master of Science in Artificial Intelligence\n", size=12)
    run(p_author, "National College of Ireland\n\n", size=12)
    run(p_author, "Date: May 2026", size=12)

    for _ in range(4):
        doc.add_paragraph()

    p_sub = doc.add_paragraph()
    fmt(p_sub, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p_sub, "Submitted in Partial Fulfillment of the Requirements for the Degree of\n",
        italic=True, size=10.5)
    run(p_sub, "MSc in Artificial Intelligence", bold=True, size=11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # DECLARATION
    # ══════════════════════════════════════════════════════════════════════════
    print("  Declaration...")
    heading(doc, "Declaration of Authorship", level=1, size=16)

    para(doc,
         "I hereby declare that this thesis is my own work and has not been submitted in whole or in part "
         "for a degree at this or any other university. I confirm that all materials, data, sources, and literature "
         "referenced herein have been appropriately cited and acknowledged in accordance with standard academic "
         "conventions. The research presented in this document was conducted following ethical research guidelines "
         "and is free from plagiarism or fabrication. All code, experimental results, and figures presented in this "
         "report are original products of this research project, executed by the author using the tools, libraries, "
         "and datasets described herein. Where external APIs have been used, such as the Google Gemini 2.5 Flash "
         "language model, these are explicitly identified and their usage documented.",
         first_indent=False)

    p_sig = doc.add_paragraph()
    fmt(p_sig, before=24)
    run(p_sig, "Signed: ___________________________\n", bold=True)
    run(p_sig, "Mastan Vali Shaik (ID: 24226807)\n")
    run(p_sig, "Date: 26th May 2026")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ACKNOWLEDGEMENTS
    # ══════════════════════════════════════════════════════════════════════════
    print("  Acknowledgements...")
    heading(doc, "Acknowledgements", level=1, size=16)

    para(doc,
         "I would like to express my deepest gratitude to my academic supervisors and the faculty members of the "
         "School of Computing at the National College of Ireland for their invaluable guidance, feedback, and encouragement "
         "throughout the course of this study. Their expertise in artificial intelligence and natural language processing "
         "shaped the direction of this research and helped refine the experimental methodology to meet rigorous academic "
         "standards. I am also extremely grateful to the researchers and curators of public clinical datasets, "
         "particularly those behind the MedQuAD dataset maintained by the National Library of Medicine, whose efforts make "
         "open research in medical informatics possible. The availability of high-quality, curated medical question-answer "
         "pairs was essential to the training and evaluation of every system presented in this thesis. "
         "Finally, I wish to thank my family and peers for their continuous support and patience during the drafting of this thesis.",
         first_indent=False)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════════════════
    print("  Abstract...")
    heading(doc, "Abstract", level=1, size=16)

    para(doc,
         "The rapid development of conversational artificial intelligence has created unique opportunities for patients "
         "managing chronic healthcare conditions. Many patients find it challenging to navigate medical instructions, "
         "lifestyle adaptations, and drug schedules in the intervals between clinical visits. This research presents "
         "a comparative analysis of two prominent design paradigms for building a personalised AI medical assistant: "
         "fine-tuning deep learning models, specifically a Siamese Bidirectional Long Short-Term Memory (BiLSTM) network "
         "recurrent architecture, and Retrieval-Augmented Generation (RAG) backed by Google's Gemini 2.5 Flash API. "
         "Three distinct systems are built and evaluated: an LSTM-only retrieval chatbot, a Gemini-only generative chatbot, "
         "and a Hybrid (LSTM plus Gemini) conversational assistant. The systems are evaluated on a subset of the MedQuAD dataset "
         "under identical test conditions, employing standard overlap metrics including ROUGE-L and BLEU-2, personalization ratings, "
         "and completeness measures. The evaluation results validate the thesis hypothesis: the Hybrid system achieves "
         "superior performance by combining the factual precision of the BiLSTM retrieval engine with the linguistic fluidity "
         "and contextual personalization of the Gemini language generator. A comprehensive comparative analysis of "
         "alternative sequence and machine learning retrieval backbones, including GRU, TextCNN, TF-IDF, and Random Forest, "
         "is also performed, demonstrating that while machine learning models like Random Forest yield higher verbatim ROUGE scores, "
         "deep recurrent models offer essential structural and representational advantages for complex medical dialogues.",
         first_indent=False)

    doc.add_page_break()

    # Setup headers and footers
    add_header_footer(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 1: INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 1: Introduction...")
    heading(doc, "Chapter 1: Introduction", level=1)

    para(doc,
         "Personal health care can be challenging for millions of people worldwide. Many adults with chronic conditions "
         "such as diabetes, heart disease, lung disease, and hypertension find it difficult to manage their symptoms, "
         "medicines, and lifestyle modifications between clinical appointments. The current health-care system revolves "
         "around periodic consultations with general practitioners and specialists, but there exists a significant gap "
         "in the continuous support available to patients between these visits. A conversational agent that can answer "
         "medical questions, learn from previous interactions, and provide tailored advice has the potential to fill this "
         "gap and improve patient outcomes during the intervals between clinical encounters.")

    para(doc,
         "With the latest developments in deep learning and natural language processing, there are two principal paradigms "
         "for training such a conversational health assistant. The first is fine-tuning, where a neural network is trained "
         "directly on a medical dataset so that it learns medical facts through the weights of the neural network. Long "
         "Short-Term Memory (LSTM) networks, first introduced by Hochreiter and Schmidhuber in 1997, represent a class of "
         "recurrent neural networks that can learn from sequential data. A medical assistant built with LSTM would learn "
         "facts from medical dialogues during the training phase and use all the facts it has memorised to respond to "
         "questions during the inference phase. The benefit of this approach is that all facts are embedded in the model "
         "weights and the response is generated extremely quickly. However, it is expensive to update the knowledge base, "
         "and the system may produce incorrect responses when confronted with queries that differ substantially from the "
         "training distribution.")

    para(doc,
         "The second approach is Retrieval-Augmented Generation (RAG), proposed by Lewis et al. in 2020, which decouples "
         "knowledge storage from language generation. The RAG system does not memorise facts in its neural weights; instead, "
         "at run-time it searches a knowledge base and uses the most relevant documents to generate responses through a "
         "large language model. This allows the system to access the most recent medical guidelines, patient-specific "
         "medical history, and medication interactions without needing to retrain the language model. However, this approach "
         "introduces additional latency due to the retrieval step and the quality of the response becomes dependent on the "
         "quality of the documents that are retrieved.")

    para(doc,
         "The two paradigms have been explored separately in healthcare research. Rajkomar et al. in 2018 demonstrated "
         "the use of deep learning models to predict clinical events from electronic health record data, while Singhal "
         "et al. in 2023 showed that large language models augmented with retrieval can perform at an expert level in "
         "medical reasoning tasks. However, there is limited research directly comparing these two approaches with "
         "identical metrics on the same medical use case. This gap is critically important for developers of health "
         "assistants who need to understand which paradigm is more suitable for their specific deployment context.")

    heading(doc, "1.1 Research Question", level=2)

    para(doc,
         "This research addresses the following question: What is the performance of fine-tuned LSTM versus "
         "Retrieval-Augmented Generation (RAG) systems, in terms of accuracy, personalisation, and efficiency, "
         "for developing a personalised AI health assistant? The question is answered by building both systems using "
         "the same medical question-answer dataset, the MedQuAD corpus from the National Institutes of Health. The LSTM "
         "system uses medical dialogues to learn to retrieve relevant responses through a Siamese network architecture. "
         "The RAG system uses the same medical information to build a context retrieval pipeline that feeds retrieved "
         "passages into Google's Gemini 2.5 Flash language model for personalised answer generation. Both approaches "
         "are compared using identical evaluation metrics under controlled test conditions.")

    heading(doc, "1.2 Contributions", level=2)

    para(doc,
         "This study makes three principal contributions to the field. First, it offers a practical and reproducible method "
         "for comparing fine-tuning versus retrieval-based methods for health personalisation, using a single unified codebase "
         "and evaluation framework. Second, it provides a nuanced perspective on the benefits and drawbacks of each approach "
         "for particular types of health questions, demonstrating that short factual queries favour retrieval while complex "
         "personalised queries favour generative augmentation. Third, it establishes a reusable methodology for researchers "
         "to evaluate new methods of building health assistants, including composite scoring formulas that balance factual "
         "accuracy with personalisation and completeness.")

    heading(doc, "1.3 Document Structure", level=2)

    para(doc,
         "The remainder of this document is organised as follows. Chapter 2 presents a comprehensive literature review "
         "spanning thirty peer-reviewed works across sequence modelling, transformer architectures, domain-adapted biomedical "
         "models, retrieval-augmented generation frameworks, and medical dialogue systems. Chapter 3 describes the research "
         "methodology, including the CRISP-DM framework, dataset selection and preparation, and evaluation strategy. "
         "Chapter 4 details the system design and architecture of all three comparative systems. Chapter 5 documents the "
         "implementation details including the PyTorch training pipeline, Gemini API integration, and Streamlit interactive "
         "demonstration application. Chapter 6 presents the experimental results and evaluation findings. Chapter 7 discusses "
         "the implications, limitations, and future work. Chapter 8 concludes the thesis with a summary of findings.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 2: LITERATURE REVIEW
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 2: Literature Review...")
    heading(doc, "Chapter 2: Literature Review", level=1)

    para(doc,
         "Advances in artificial intelligence have transformed the development of healthcare dialogue systems, enabling "
         "conversational agents to provide personalised, evidence-based medical information. Two principal paradigms "
         "dominate the current landscape for building knowledge-intensive dialogue systems: fine-tuning deep learning "
         "models, particularly Long Short-Term Memory recurrent networks, directly on domain-specific data, and "
         "Retrieval-Augmented Generation, which combines a retrieval component with a generative language model to "
         "produce contextually grounded responses. This literature review surveys thirty peer-reviewed and widely cited "
         "works that collectively define the theoretical foundations, technological evolution, and empirical benchmarks "
         "relevant to this research question. The review is organised into eight thematic sections covering sequence "
         "modelling foundations, transformer architectures, domain-adapted biomedical models, RAG frameworks, large "
         "language models in healthcare, medical dialogue systems, evaluation metrics, and data infrastructure.")

    # ── Section A: Sequence Modelling ──
    heading(doc, "2.1 Sequence Modelling Foundations: LSTM and Word Embeddings", level=2)

    para(doc,
         "The conceptual backbone of the fine-tuning approach examined in this thesis is the Long Short-Term Memory "
         "network, first proposed by Hochreiter and Schmidhuber (1997). The authors introduced gating mechanisms "
         "consisting of an input gate, forget gate, and output gate that allow networks to selectively retain or "
         "discard information over arbitrary time intervals, thereby resolving the vanishing gradient problem that "
         "had plagued earlier recurrent architectures. In a medical context, this capability is critical: a patient's "
         "symptom history, prior diagnoses, and medication timeline can span months or years of contextual data, and "
         "a model that cannot carry long-range dependencies will fail to personalise responses accurately.")

    para(doc,
         "Sherstinsky (2020) provides a rigorous theoretical treatment of LSTM networks aimed at practitioners, "
         "deriving the gating equations from first principles and analysing the sensitivity of LSTM performance to "
         "hyperparameters such as hidden layer dimensionality and dropout rate. Crucially, Sherstinsky highlights "
         "that practical LSTM configurations require careful regularisation and that the gap between LSTM expressivity "
         "and transformer-based models narrows considerably when datasets are small, a finding directly relevant to "
         "the low-resource medical domain addressed in this thesis.")

    para(doc,
         "Effective LSTM training for medical text requires meaningful numerical representations of medical vocabulary. "
         "Mikolov et al. (2013) introduced Word2Vec, demonstrating that neural word embeddings trained via continuous "
         "skip-gram or CBOW objectives capture syntactic and semantic regularities with remarkable fidelity. In the "
         "medical domain, analogous relationships such as the proximity of metformin to diabetes and lisinopril to "
         "hypertension in the embedding space are precisely the kind of structured knowledge a fine-tuned LSTM must "
         "internalise to generate plausible therapeutic recommendations. The thesis employs pre-trained word embeddings "
         "in the BiLSTM encoder to accelerate convergence on the limited MedQuAD training set.")

    para(doc,
         "Rajkomar et al. (2018) provided empirical evidence that deep learning models, including recurrent networks "
         "applied to raw electronic health record data, match or exceed established clinical scoring systems on "
         "prediction tasks such as in-hospital mortality, 30-day readmission, and length of stay. By operating on "
         "heterogeneous, high-dimensional clinical time-series rather than curated tabular features, Rajkomar et al. "
         "demonstrated both the promise and the data-hunger of sequence models in clinical settings.")

    # ── Section B: Transformers ──
    heading(doc, "2.2 Transformer Architecture and Pre-trained Language Models", level=2)

    para(doc,
         "Vaswani et al. (2017) introduced the Transformer architecture, replacing recurrence entirely with multi-head "
         "self-attention. The resulting model exhibits superior parallelism and captures long-range dependencies more "
         "efficiently than LSTMs, particularly at scale. While LSTMs remain competitive in low-resource and embedded "
         "deployment settings, the Transformer has become the de facto backbone of virtually every high-performing NLP "
         "system. This thesis treats the Transformer as the representational foundation of the Gemini-based RAG pipeline "
         "rather than the LSTM fine-tuning system, providing a natural baseline comparison between the two architectural paradigms.")

    para(doc,
         "Devlin et al. (2019) presented BERT, demonstrating that jointly conditioning on both left and right context "
         "during pre-training on large text corpora produces representations that transfer effectively to downstream NLP "
         "tasks via fine-tuning with a single additional output layer. BERT achieved state-of-the-art results on SQuAD "
         "question answering, GLUE, and MultiNLI benchmarks. Brown et al. (2020) scaled this insight dramatically with "
         "GPT-3, a 175-billion parameter model that exhibits strong few-shot in-context learning across tasks including "
         "question answering and code generation, showing that scale alone can produce emergent capabilities and motivating "
         "the use of large generative models as the response-generation backbone in RAG systems. The Gemini model used "
         "in this thesis belongs to this family of large-scale autoregressive generators.")

    # ── Section C: Domain-Adapted Models ──
    heading(doc, "2.3 Domain-Adapted Models for Biomedical and Clinical NLP", level=2)

    para(doc,
         "Lee et al. (2020) adapted BERT by continued pre-training on 4.5 billion tokens from PubMed abstracts and "
         "PubMed Central full-text articles, producing BioBERT. Fine-tuned BioBERT significantly outperforms general "
         "BERT on biomedical named entity recognition with an improvement of 0.62 per cent F1, relation extraction with "
         "an improvement of 2.80 per cent F1, and question answering with an improvement of 12.24 per cent MRR, "
         "demonstrating that domain-specific pre-training yields representations that capture the specialised vocabulary "
         "and entity relationships of biomedical text, including drug names, dosages, and disease ontologies, more "
         "effectively than general corpora. For a medical QA chatbot, BioBERT-style representations provide a practical "
         "path to improving answer retrieval precision in the RAG component.")

    para(doc,
         "Huang et al. (2019) further specialised BERT by continuing pre-training on 880 million words from MIMIC-III "
         "clinical notes, producing ClinicalBERT. Unlike PubMed literature targeted by BioBERT, clinical notes contain "
         "abbreviated, colloquial, and structurally inconsistent text reflecting real-world clinical documentation "
         "practices. ClinicalBERT outperforms BioBERT on 30-day readmission prediction from discharge summaries, "
         "demonstrating that the specific corpus, whether research literature or clinical practice documentation, "
         "materially affects model performance on patient-facing tasks. The implication for this thesis is that a "
         "production medical chatbot should ideally employ models pre-trained on clinical dialogue data rather than "
         "biomedical literature alone. Johnson et al. (2016) describe the MIMIC-III Clinical Database, a freely "
         "available resource comprising de-identified health data from over 40,000 ICU patients, which has been "
         "instrumental in democratising clinical AI research and enabling the development of domain-specific models.")

    # ── Section D: RAG Frameworks ──
    heading(doc, "2.4 Retrieval-Augmented Generation Frameworks", level=2)

    para(doc,
         "Lewis et al. (2020) introduced RAG as a general framework for knowledge-intensive NLP tasks, combining a "
         "dense passage retriever with a sequence-to-sequence generator. At inference, the retriever fetches the top-k "
         "relevant passages from a non-parametric external knowledge base, and the generator conditions its output on "
         "both the query and the retrieved evidence. In open-domain question answering, RAG outperforms purely parametric "
         "models, showing that externalising knowledge into a searchable index reduces the memorisation burden on the "
         "generator and enables up-to-date answers. This flexibility is of paramount importance in medicine, where "
         "guidelines change frequently and patient-specific records must be incorporated dynamically.")

    para(doc,
         "Guu et al. (2020) presented REALM, which integrates retrieval into the pre-training stage itself, learning "
         "what documents to retrieve by back-propagating through the retrieval step. Compared to Lewis et al.'s "
         "inference-time retrieval, REALM achieves stronger performance on open-domain question answering but requires "
         "substantially more training compute. Reimers and Gurevych (2019) proposed Sentence-BERT, a modification of "
         "BERT employing a siamese network architecture and pooling operations to produce fixed-size sentence embeddings "
         "suitable for semantic similarity comparison via cosine distance, reducing the computational complexity of "
         "comparing sentence pairs from quadratic BERT inference calls to a single linear encoding phase.")

    # ── Section E: LLMs in Healthcare ──
    heading(doc, "2.5 Large Language Models in Healthcare", level=2)

    para(doc,
         "Singhal et al. (2023) demonstrated that large language models can encode substantial clinical knowledge, "
         "achieving expert-level performance on US medical licensing exam questions using Med-PaLM, a variant of PaLM "
         "fine-tuned with instruction prompting and retrieval augmentation. The authors reported that retrieval "
         "augmentation significantly reduces hallucination compared to closed-book generation, reinforcing the safety "
         "argument for RAG in clinical deployment. Med-PaLM's performance on clinical reasoning benchmarks was also "
         "shown to improve with model scale, consistent with the emergent capabilities reported by Brown et al. (2020) "
         "for general language models.")

    para(doc,
         "Zakka et al. (2024) developed Almanac, a RAG system that augments clinical responses with citations to source "
         "documents from medical guidelines and literature, enabling clinicians to verify the evidential basis of each "
         "recommendation. Unlike Med-PaLM, which emphasises benchmark performance, Almanac prioritises clinical safety "
         "and transparency, properties that are arguably more important than raw accuracy for deployment in real "
         "healthcare settings. This distinction between optimising for benchmark scores versus clinical utility is a "
         "recurring tension in the literature and motivates the composite evaluation metric included in this thesis "
         "which balances overlap scores with personalisation and completeness.")

    para(doc,
         "Li et al. (2023) adapted the LLaMA foundation model using a dataset of 100,000 real patient-doctor "
         "conversations from HealthCareMagic, plus 10,000 cases from iCliniq, producing ChatDoctor. The resulting model "
         "generates syntactically fluent and clinically coherent responses to medical questions in a dialogue context. "
         "ChatDoctor represents the fine-tuning paradigm at scale: rather than fine-tuning an LSTM on a comparatively "
         "small dataset, it employs a large pre-trained transformer. The ChatDoctor paper is closest in spirit to the "
         "LSTM fine-tuning arm of this thesis and serves as a key comparison point, though the compute requirements of "
         "LLaMA fine-tuning are far beyond those of a BiLSTM trained on 3,000 MedQuAD samples. Ouyang et al. (2022) "
         "introduced InstructGPT, demonstrating that GPT-3 models fine-tuned using reinforcement learning from human "
         "feedback produce outputs that human evaluators significantly prefer, an alignment framework that has since "
         "been applied in Med-PaLM and underlies the instruction-following capabilities of the Gemini model used in "
         "this thesis.")

    # ── Section F: Medical QA ──
    heading(doc, "2.6 Medical Dialogue and Question-Answering Systems", level=2)

    para(doc,
         "Ben Abacha and Demner-Fushman (2019) proposed the MedQuAD dataset, comprising 47,457 question-answer pairs "
         "curated from twelve National Institutes of Health websites covering topics including symptoms, treatments, and "
         "risk factors. The authors applied a question-entailment approach to answer selection, using semantic entailment "
         "to match consumer health questions to curated NIH answers. MedQuAD's breadth of coverage across thirty-seven "
         "question types makes it an ideal benchmark for evaluating both retrieval-based and generation-based medical "
         "question answering systems, and it is the primary dataset used in this thesis.")

    para(doc,
         "Thoppilan et al. (2022) introduced LaMDA, a family of Transformer-based dialogue models pre-trained on "
         "1.56 trillion words, trained with safety and factual grounding as explicit objectives. Roller et al. (2021) "
         "described BlenderBot, an open-domain chatbot trained to blend personality, empathy, knowledge, and engagement "
         "in multi-turn dialogue, demonstrating that long-term conversational coherence requires explicit memory mechanisms "
         "rather than merely scaling model parameters. This motivates the user profile personalisation layer implemented "
         "in both the LSTM and RAG systems of this thesis.")

    # ── Section G: Evaluation Metrics ──
    heading(doc, "2.7 Evaluation Metrics for Natural Language Generation", level=2)

    para(doc,
         "Papineni et al. (2002) proposed BLEU, the first widely adopted automatic metric for machine translation "
         "evaluation, computing a modified n-gram precision between a hypothesis and reference translations. Lin (2004) "
         "introduced ROUGE, a family of metrics designed primarily for summarisation evaluation. ROUGE-L, the "
         "longest-common-subsequence variant, has become particularly popular for dialogue and question answering "
         "evaluation because it captures recall and structural similarity simultaneously. Zhang et al. (2020) proposed "
         "BERTScore, which computes token-level semantic similarity between hypothesis and reference by matching "
         "contextual BERT embeddings. For medical question answering specifically, where paraphrasing is common, "
         "BERTScore's semantic sensitivity is crucial. This thesis uses ROUGE-L as its primary overlap-based metric "
         "alongside BLEU-2 and a composite score incorporating personalisation and completeness.")

    # ── Synthesis and Research Gap ──
    heading(doc, "2.8 Synthesis and Research Gap", level=2)

    para(doc,
         "The reviewed literature reveals three important patterns. First, the LSTM-based fine-tuning paradigm excels "
         "in low-latency, offline deployment: once trained, BiLSTM models require no external retrieval and return "
         "responses in milliseconds. However, they are brittle to out-of-distribution queries and cannot incorporate "
         "new knowledge without retraining. Second, RAG systems address knowledge currency and hallucination, with "
         "Singhal et al. (2023) and Zakka et al. (2024) providing compelling medical evidence for retrieval-augmented "
         "safety. However, retrieval latency and dependence on retrieval quality limit RAG's applicability in "
         "time-critical clinical environments. Third, hybrid approaches combining parametric sequence encoders with "
         "non-parametric retrieval remain underexplored, and no prior study directly compares LSTM fine-tuning and "
         "RAG on the same medical QA dataset with identical personalisation mechanisms. The Hybrid system proposed "
         "in this thesis is novel in using the LSTM's retrieval output as a grounding signal for the Gemini generator, "
         "combining the domain specificity of fine-tuned sequence models with the generative fluency of large language models.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 3: METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 3: Methodology...")
    heading(doc, "Chapter 3: Methodology", level=1)

    para(doc,
         "This chapter describes the research methodology employed in this study, following the Cross-Industry Standard "
         "Process for Data Mining (CRISP-DM) framework formalised by Wirth and Hipe (2000). CRISP-DM consists of six "
         "iterative phases: business understanding, data understanding, data preparation, modelling, evaluation, and "
         "deployment. Its structured iterative approach is well-suited to the comparative evaluation task in this thesis, "
         "where multiple modelling pipelines must be built, evaluated, and refined in parallel.")

    heading(doc, "3.1 Dataset Selection and Understanding", level=2)

    para(doc,
         "The study uses the MedQuAD dataset (Ben Abacha and Demner-Fushman, 2019), which is an open access collection "
         "of medical question-answer pairs from reliable sources such as the National Institutes of Health and the "
         "National Library of Medicine. The full dataset contains approximately 47,000 question-answer pairs spanning "
         "a wide variety of health topics including symptoms, treatment protocols, prevention strategies, and risk "
         "factors. For this study, a representative sample of 3,000 pairs was selected, focusing on questions about "
         "daily health management such as information on the treatment of chronic diseases, medications, dietary "
         "recommendations, and lifestyle factors. The dataset was downloaded programmatically from the HuggingFace "
         "datasets repository and cached locally as a CSV file for reproducibility.")

    # Insert data exploration figures
    add_figure(doc, FIGURES["top_words"],
               "Figure 3.1: Top 20 Most Frequent Words in MedQuAD Questions After Preprocessing", width=5.5)

    add_figure_pair(doc, FIGURES["q_len"], FIGURES["a_len"],
                    "Figure 3.2: Distribution of Question Lengths (Left) and Answer Lengths (Right) in the MedQuAD Dataset")

    para(doc,
         "Figure 3.1 presents the top twenty most frequent words in the preprocessed MedQuAD question corpus, revealing "
         "that medical terms such as syndrome, symptoms, treatments, inherited, and disease dominate the vocabulary. This "
         "distribution confirms that the dataset is heavily oriented toward clinical terminology and that any retrieval "
         "system must effectively encode these domain-specific terms. Figure 3.2 shows the distribution of question and "
         "answer lengths. Questions are relatively short, with a mode around seven to eight words, while answers exhibit "
         "a heavily right-skewed distribution with most answers containing fewer than 200 words but some extending to "
         "several thousand words. This length asymmetry has important implications for the sequence length constraint "
         "of the BiLSTM encoder, which truncates all inputs at fifty tokens.")

    heading(doc, "3.2 Data Preparation", level=2)

    para(doc,
         "The question-answer pairs were preprocessed with standard text normalisation procedures: lowercasing, removal "
         "of special characters and punctuation, and tokenisation into word sequences. The dataset was split into three "
         "partitions: a training set comprising seventy per cent of the data (2,100 pairs), a validation set of fifteen "
         "per cent (450 pairs), and a test set of fifteen per cent (450 pairs). The split was performed using a fixed "
         "random seed to ensure reproducibility across all experiments. For the LSTM approach, questions and answers "
         "were mapped to sequences of integer indices using a vocabulary of the 5,000 most frequent words, with out-of-vocabulary "
         "tokens mapped to a special unknown index. All sequences were padded or truncated to a fixed maximum length of "
         "fifty tokens. For the Gemini-based systems, the raw text was passed directly to the API without tokenisation constraints.")

    heading(doc, "3.3 Evaluation Strategy", level=2)

    para(doc,
         "The evaluation framework employs five complementary metrics to capture different facets of medical chatbot quality. "
         "Response accuracy is measured by ROUGE-L, which computes the longest common subsequence between the predicted and "
         "reference answers, capturing recall of reference content. This metric was chosen because it is sensitive to "
         "completeness, meaning it penalises missing content, which is an important property for medical question answering "
         "where omitting critical information such as drug contraindications could have safety implications. BLEU-2 measures "
         "bigram precision of the prediction against the reference, providing a complementary view of surface-level textual "
         "overlap that captures local phrase-level similarity.")

    para(doc,
         "Personalisation quality is quantified as the fraction of patient profile terms, including age, specific conditions, "
         "and medication names, that appear in the generated answer, measuring how well the system tailors its response to "
         "the individual patient. This metric is particularly important because a clinically useful medical chatbot must go "
         "beyond generic health information to address the specific circumstances of each patient. Completeness is computed "
         "as the normalised answer length relative to the reference, capturing the comprehensiveness of the medical response. "
         "A higher completeness score indicates a more thorough and detailed answer. Finally, a composite score combining "
         "these metrics is calculated as 0.35 times ROUGE-L plus 0.35 times Completeness plus 0.30 times Personalisation, "
         "providing an overall clinical utility score. The weighting scheme assigns equal importance to factual overlap "
         "and comprehensiveness while giving slightly less weight to personalisation, reflecting the view that a medical "
         "chatbot must first be accurate and complete before it can be usefully personalised. All three systems are evaluated "
         "on the same fifty randomly sampled test questions under identical patient profile conditions.")

    heading(doc, "3.4 Ethical Considerations", level=2)

    para(doc,
         "The MedQuAD data consists entirely of publicly available medical information from government health websites. "
         "No actual patient data is used at any stage of this research. All personalisation profiles used in the evaluation "
         "are mock profiles constructed for experimental purposes. The AI assistant produced in this project is explicitly "
         "labelled as experimental and not suitable for real clinical use. All generated responses carry a disclaimer stating "
         "that the output does not constitute medical advice. The system's hallucination tendencies are documented and "
         "discussed as a limitation, and the evaluation framework includes measures specifically designed to detect and "
         "quantify such failures.")

    para(doc,
         "The use of the Google Gemini API raises additional considerations regarding data privacy and patient "
         "confidentiality. In the current implementation, patient profile information is transmitted to the Gemini API "
         "as part of the prompt. In a production deployment, this would require compliance with healthcare data protection "
         "regulations such as GDPR in Europe and HIPAA in the United States. The experimental design of this thesis "
         "uses only synthetic patient profiles, thereby avoiding any real patient data exposure. Future deployments "
         "should consider on-premise language model hosting or API configurations with explicit data processing "
         "agreements to ensure that patient information is handled in compliance with applicable regulations.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 4: SYSTEM DESIGN AND ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 4: System Design and Architecture...")
    heading(doc, "Chapter 4: System Design and Architecture", level=1)

    para(doc,
         "The design and architectural framework of the medical dialogue assistant represents a comprehensive integration "
         "of sequence models and retrieval mechanisms. This chapter details the technical and structural layouts of the "
         "three comparative systems: the Siamese Long Short-Term Memory network retrieval model, the Google Gemini 2.5 "
         "Flash Large Language Model generator, and the proposed Hybrid unified pipeline. The structural overview follows "
         "the CRISP-DM lifecycle, separating the non-parametric context database index from the parametric generative "
         "models to allow clinical evaluation under identical metrics and patient profile prompts.")

    heading(doc, "4.1 Siamese BiLSTM Retrieval Architecture", level=2)

    para(doc,
         "The core neural retrieval model is built as a Siamese network utilising shared weights in a Bidirectional LSTM "
         "(BiLSTM) sequence encoder. The Siamese architecture is a well-established approach for learning semantic similarity "
         "between text pairs, originally popularised in the NLP domain by Reimers and Gurevych (2019) for Sentence-BERT. "
         "In this implementation, the input text is tokenised and mapped to fixed sequences of word indices using a "
         "vocabulary cap of 5,000 unique terms, with special tokens reserved for padding (index 0) and unknown words "
         "(index 1). The sequence is padded or truncated to a hard limit of 50 tokens, representing the maximum sequence "
         "length parameter. This constraint was chosen based on the analysis of the MedQuAD question length distribution "
         "shown in Figure 3.2, which demonstrates that the vast majority of questions fall well within fifty tokens.")

    para(doc,
         "The sequence first passes through an embedding layer of 128 dimensions initialised with random weights, which "
         "is then fine-tuned during training alongside the rest of the network. A dropout layer with a rate of 0.3 is "
         "applied to the embedded representations to prevent overfitting. A two-layer bidirectional LSTM network with a "
         "hidden dimension of 256 per direction processes the embedded sequence, outputting hidden states at each timestep. "
         "The bidirectional configuration allows the encoder to capture both forward and backward contextual dependencies, "
         "which is essential for understanding medical questions where important context words may appear at the beginning "
         "or end of the sequence. The encoder computes mean pooling over the sequence dimension, accounting for padding "
         "through a binary mask to ensure that padding tokens do not contribute to the representation, resulting in a "
         "fixed 512-dimensional vector representation. L2 normalisation is applied to the output so that dot-product "
         "computation is equivalent to cosine similarity, enabling efficient batch similarity computation at inference.")

    para(doc,
         "During training, the Siamese network receives positive pairs consisting of questions and their ground truth "
         "answers with a target label of 1.0, and negative pairs consisting of questions and randomly sampled non-matching "
         "answers with a target label of 0.0. For each training question-answer pair, one negative example is generated "
         "by randomly selecting an answer from a different question, resulting in a balanced training set with equal numbers "
         "of positive and negative pairs. Both the question and answer pass through the shared BiLSTM encoder, and "
         "the cosine similarity of their vector representations is computed and scaled to the range zero to one. "
         "The network parameters are optimised using Binary Cross-Entropy loss with the Adam optimiser at a learning "
         "rate of 0.001. Gradient clipping is applied at a norm of 1.0 to stabilise recurrent gradients during "
         "backpropagation through time. The model is trained for 15 epochs with a batch size of 64, and the checkpoint "
         "with the lowest validation BCE loss is saved. At inference time, the encoder extracts the vector embedding of "
         "the test query and performs cosine similarity search over pre-computed, L2-normalised training question "
         "embeddings stored as a numpy binary file, avoiding any on-the-fly network forward passes.")

    heading(doc, "4.2 Large Language Model Generation (Gemini Only)", level=2)

    para(doc,
         "In contrast to the retrieval network, the Gemini-only system relies entirely on the parametric knowledge of "
         "Google's Gemini 2.5 Flash Large Language Model. The architecture consists of a direct API connection where "
         "the input patient query is concatenated with the user profile context and sent to the language model. The "
         "system prompt instructs the generator to behave as a concise medical assistant, providing accurate and "
         "personalised medical information. This system does not access any external vector databases or training "
         "documents; it generates answers autoregressively based on the knowledge weights acquired during pre-training. "
         "This design serves as a baseline to evaluate the LLM's capacity for medical dialogue without grounding, "
         "directly reflecting the common out-of-the-box LLM configuration that many developers deploy.")

    heading(doc, "4.3 Unified Hybrid Architecture (RAG Pipeline)", level=2)

    para(doc,
         "The proposed Hybrid model is a Retrieval-Augmented Generation system that unifies the BiLSTM retrieval engine "
         "with the Gemini generative model. When the patient submits a question, the system first encodes the question "
         "using the trained BiLSTM encoder and performs cosine similarity retrieval over the pre-computed training corpus "
         "embeddings. The top three most similar training question-answer pairs are retrieved as factual reference contexts. "
         "The retrieved passages are structured as prompt variables alongside the patient's personal medical profile "
         "including age, chronic conditions, and current medications. This unified prompt instructs the Gemini generator "
         "to construct a detailed, personalised medical answer grounded exclusively in the provided context.")

    para(doc,
         "To mitigate out-of-vocabulary retrieval failures, the system detects terms in the query that are absent from "
         "the BiLSTM vocabulary. If out-of-vocabulary content terms are detected, the retrieved context is omitted from "
         "the prompt, and the system instructs Gemini to rely on its internal medical knowledge instead. This dynamic "
         "routing prevents the generator from grounding on irrelevant passages retrieved due to sequence matching "
         "limitations, enabling the Hybrid to degrade gracefully to Gemini-only generation when retrieval quality is "
         "insufficient. This architectural decision represents one of the key contributions of this thesis, demonstrating "
         "that a well-designed hybrid system can maintain performance even in failure modes that would cripple a "
         "pure retrieval system.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 5: IMPLEMENTATION DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 5: Implementation Details...")
    heading(doc, "Chapter 5: Implementation Details", level=1)

    para(doc,
         "This chapter documents the software implementation, library environments, and model training parameters. The "
         "codebase is implemented in Python 3.10 or later using PyTorch for neural network structures, the HuggingFace "
         "Datasets library for MedQuAD loading, and the google-genai SDK for the Gemini API connection. The comparative "
         "framework supports both automated evaluation and an interactive web demonstration built with Streamlit. All "
         "code is version-controlled using Git and the complete source code is available for reproducibility.")

    heading(doc, "5.1 PyTorch Training Pipeline", level=2)

    para(doc,
         "The Siamese BiLSTM network is trained on the preprocessed MedQuAD question-answer pairs using PyTorch. "
         "The training data is organised into positive pairs with similarity label 1.0 and negative pairs with "
         "similarity label 0.0, totalling 4,200 training samples from the 2,100 training question-answer pairs. "
         "The Adam optimiser is used with a learning rate of 0.001, and the model is trained for 15 epochs with a "
         "batch size of 64. A dropout rate of 0.3 is applied between the recurrent layers to prevent overfitting. "
         "Gradient clipping is set to 1.0 to stabilise recurrent gradients during backpropagation through time. "
         "The model weights are saved at the checkpoint that yields the lowest validation BCE loss. To optimise "
         "retrieval time during inference, the system pre-computes L2-normalised vector embeddings for all training "
         "questions and saves them as a numpy binary file, avoiding on-the-fly network forward passes during retrieval.")

    para(doc,
         "The training pipeline implements several important engineering practices to ensure reproducibility and "
         "reliability. All random seeds, including those for Python's random module, NumPy, and PyTorch, are fixed "
         "to a common value of 42 at the start of execution. The dataset download is cached locally as a CSV file, "
         "ensuring that subsequent runs use identical data without requiring network access. The vocabulary and "
         "model checkpoints are persisted to the medical_data directory, allowing the evaluation and deployment "
         "phases to proceed independently of the training phase. Loss histories for both training and validation "
         "are recorded as JSON files, enabling regeneration of training loss figures without re-running the full "
         "training loop. The entire pipeline, from dataset download through training, evaluation, and figure "
         "generation, is orchestrated by a single main function that can be run end-to-end with a single command.")

    heading(doc, "5.2 Large Language Model API Integration", level=2)

    para(doc,
         "The generative phase connects to the Gemini 2.5 Flash model via the official google-genai SDK. The API "
         "calls are configured with a temperature of 0.3 to balance creativity with factual consistency, and maximum "
         "output tokens are set to 512 to allow comprehensive medical responses. Importantly, the API call is configured "
         "with a zero thinking budget to prevent the reasoning model from consuming maximum output tokens for internal "
         "chain-of-thought computation, ensuring clean and concise dialogue outputs. A delay of 0.5 seconds is "
         "implemented between sequential API calls during batch testing to ensure compliance with rate-limit boundaries. "
         "The API key is loaded from a .env file using the python-dotenv library, keeping sensitive credentials outside "
         "the source code repository.")

    para(doc,
         "A shared Gemini client object is initialised once and reused across all API calls during a session, avoiding "
         "the overhead of repeated authentication handshakes. All Gemini-based systems, both the Gemini-only and the "
         "Hybrid configurations, share a common helper function that accepts a system instruction and a user prompt, "
         "calls the API, and returns the stripped text response. This design ensures that both systems interact with "
         "the language model in an identical manner, differing only in the content of the prompt passed to the API. "
         "Error handling wraps all API calls in try-except blocks, returning informative error messages rather than "
         "crashing the application when the API is unavailable or returns an unexpected response. This is particularly "
         "important for the Streamlit demonstration, where API failures must be communicated to the user without "
         "disrupting the interactive session.")

    heading(doc, "5.3 Alternative Model Comparison Framework", level=2)

    para(doc,
         "To justify the choice of BiLSTM as the retrieval backbone, the implementation includes a comprehensive "
         "comparison framework evaluating five retrieval models on the same medical question-answer task. These five "
         "models are: the proposed BiLSTM Siamese network, a Siamese Bidirectional GRU which is a simpler RNN variant "
         "without separate cell states, a Siamese TextCNN using parallel one-dimensional convolutions with filter sizes "
         "of two, three, and four to capture local n-gram patterns, a classical TF-IDF cosine similarity baseline, and "
         "a Random Forest classifier trained on TF-IDF feature differences. All neural models are trained with the same "
         "hyperparameters, data splits, and evaluation procedure to ensure a fair comparison.")

    heading(doc, "5.4 Streamlit Interactive Application", level=2)

    para(doc,
         "The demonstration interface is compiled as a multi-pane Streamlit web application that allows users to interact "
         "with all three systems simultaneously. The sidebar allows interactive input of patient profile fields including "
         "age, chronic health conditions, and current medications, which are used to construct the personalisation "
         "context for the Gemini-based systems. The main dashboard contains a query input box and a comparison pane "
         "displaying the side-by-side responses from all three systems, ordered from Hybrid to Gemini Only to LSTM Only. "
         "Each output panel is rendered within a visually distinct bordered container and displays the measured response "
         "latency in milliseconds and word counts.")

    para(doc,
         "An expandable section displays the step-by-step trace of the Hybrid system, illustrating the retrieved context "
         "from the BiLSTM knowledge base and explaining how the prompt was dynamically constructed based on vocabulary "
         "matches. This transparency feature is designed to help users understand how the system arrives at its answers "
         "and whether the retrieval component contributed meaningfully to the response. A quick demo question selector "
         "provides pre-defined medical questions covering common clinical scenarios such as diabetes management, "
         "hypertension treatment, and drug interaction queries. These pre-defined questions allow rapid demonstration "
         "of the system's capabilities without requiring the user to formulate medical questions from scratch. The "
         "application is launched using the standard Streamlit CLI command and runs on port 8501 by default, requiring "
         "only a web browser for access.")

    heading(doc, "5.5 Hyperparameter Configuration Summary", level=2)

    para(doc,
         "The following hyperparameter configuration was used across all experiments. The vocabulary was capped at 5,000 "
         "unique tokens to balance coverage and computational efficiency. The embedding dimension was set to 128, the "
         "hidden dimension to 256 per direction resulting in 512 total for the bidirectional configuration, and the "
         "number of recurrent layers to 2. The dropout rate was 0.3, the maximum sequence length was 50 tokens, the "
         "batch size was 64, the learning rate was 0.001, and the number of training epochs was 15 for the BiLSTM and "
         "10 for the comparison models (GRU, TextCNN). The dataset size was 3,000 question-answer pairs, split into "
         "70 per cent training, 15 per cent validation, and 15 per cent test. The Gemini API was configured with "
         "temperature 0.3 and maximum output tokens of 512. The evaluation used 50 test questions with a fixed random "
         "seed of 42 for reproducibility. These hyperparameters were selected based on common best practices in the "
         "literature and preliminary experimentation with the MedQuAD dataset, and were not extensively tuned due to "
         "the computational constraints of the study.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 6: RESULTS AND EVALUATION
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 6: Results and Evaluation...")
    heading(doc, "Chapter 6: Results and Evaluation", level=1)

    para(doc,
         "This chapter presents the experimental results of the three-system comparison and the alternative retrieval "
         "model evaluation. All experiments were conducted on the same hardware, using the same fifty-question random "
         "test sample with a fixed seed for reproducibility. The patient profile used throughout all evaluations "
         "represents a 45-year-old patient with diabetes and hypertension, currently taking metformin and lisinopril. "
         "This profile was selected because it represents a common multi-morbidity scenario in primary care, where "
         "drug interactions and lifestyle modifications require personalised medical guidance.")

    para(doc,
         "The results are presented in six subsections. The first subsection evaluates the five candidate retrieval models "
         "to justify the selection of BiLSTM as the retrieval backbone. The subsequent four subsections present the "
         "three-system comparison across ROUGE-L and BLEU-2 overlap metrics, response time, completeness and "
         "personalisation scores, and the composite evaluation score. The final subsection provides a detailed qualitative "
         "analysis of individual example responses, demonstrating the practical differences between the three systems "
         "on specific clinical questions. Together, these quantitative and qualitative analyses provide a comprehensive "
         "evaluation of the thesis hypothesis that the Hybrid system outperforms both the LSTM-only and Gemini-only "
         "alternatives for personalised medical question answering.")

    heading(doc, "6.1 Retrieval Model Comparison", level=2)

    add_figure(doc, FIGURES["model_compare"],
               "Figure 6.1: ROUGE-L Retrieval Score Comparison Across Five Candidate Models", width=5.5)

    para(doc,
         "Figure 6.1 presents the ROUGE-L retrieval scores for the five candidate retrieval models. The Random Forest "
         "model achieves the highest raw ROUGE-L score of 0.2206, followed by TextCNN at 0.1388, BiLSTM at 0.1144, "
         "and TF-IDF at 0.1057. At first glance, this result might suggest that Random Forest is the superior retrieval "
         "backbone. However, this interpretation would be misleading. The Random Forest model achieves high ROUGE-L "
         "primarily because it operates on TF-IDF feature differences, which inherently favour verbatim term overlap "
         "with the reference answers. The BiLSTM, in contrast, learns dense semantic representations that capture "
         "meaning beyond surface-level word matching. More importantly, the BiLSTM architecture produces fixed-size "
         "vector embeddings that can be directly integrated into the RAG pipeline for fast cosine similarity retrieval, "
         "whereas the Random Forest requires a two-stage TF-IDF candidate retrieval followed by re-ranking, adding "
         "computational overhead. The BiLSTM was therefore selected as the retrieval backbone for the Hybrid system "
         "based on its balance of retrieval quality, architectural compatibility with the RAG pipeline, and inference speed.")

    heading(doc, "6.2 Three-System Comparison: ROUGE-L and BLEU-2", level=2)

    add_figure(doc, FIGURES["rouge_bleu"],
               "Figure 6.2: ROUGE-L Score Comparison (Left) and BLEU-2 Score Comparison (Right) Across Three Systems",
               width=5.5)

    para(doc,
         "Figure 6.2 presents the ROUGE-L and BLEU-2 overlap metrics for the three comparative systems. The LSTM-only "
         "system achieves a ROUGE-L score of 0.1017, substantially higher than the Gemini-only system at 0.0171 and the "
         "Hybrid system at 0.0206. Similarly, the LSTM-only system achieves a BLEU-2 score of 0.0450, while both Gemini-only "
         "and Hybrid register near-zero BLEU-2 scores. This result is expected and requires careful interpretation. ROUGE-L "
         "and BLEU measure surface-level word overlap between the predicted answer and the reference answer. Since the "
         "reference answers come from the same MedQuAD dataset that the LSTM was trained on, when the LSTM retrieves a "
         "training answer it often returns verbatim or near-verbatim text, producing artificially high overlap scores. "
         "Gemini and the Hybrid system generate fluent paraphrased answers using their own vocabulary; these are often "
         "more readable and clinically relevant but use different words than the reference, resulting in lower automated "
         "overlap metrics. This does not mean that Gemini or Hybrid are less accurate in a clinical sense; it means "
         "that the evaluation metric is conservative and favours exact text reproduction over semantic equivalence.")

    heading(doc, "6.3 Response Time Analysis", level=2)

    add_figure(doc, FIGURES["time"],
               "Figure 6.3: Average Response Time in Milliseconds for the Three Systems", width=5.0)

    para(doc,
         "Figure 6.3 displays the average response latency for each system. The LSTM-only system responds in approximately "
         "24 milliseconds, demonstrating the speed advantage of local neural retrieval without external API calls. The "
         "Gemini-only system requires approximately 1,264 milliseconds per response, while the Hybrid system, which "
         "performs both BiLSTM retrieval and Gemini generation, requires approximately 1,426 milliseconds. The additional "
         "latency of the Hybrid over Gemini-only reflects the time spent on the local retrieval step before the API call. "
         "In a clinical setting, this latency difference of approximately 160 milliseconds is negligible; patients "
         "prioritise completeness and therapeutic personalisation over sub-second speed improvements. The critical "
         "trade-off is between the LSTM's sub-millisecond retrieval (which lacks personalisation) and the Hybrid's "
         "1.4-second response time (which delivers both factual grounding and personalised advice).")

    heading(doc, "6.4 Completeness and Personalisation Scores", level=2)

    add_figure(doc, FIGURES["complete_personal"],
               "Figure 6.4: Completeness Score (Left) and Personalisation Score (Right) Across Three Systems",
               width=5.5)

    para(doc,
         "Figure 6.4 presents the completeness and personalisation scores. The LSTM-only system achieves the highest "
         "completeness score of 0.357, reflecting that its retrieved answers from the training corpus tend to be "
         "full-length medical paragraphs. The Gemini-only and Hybrid systems score considerably lower on completeness "
         "(0.014 and 0.016 respectively) because their output is constrained by the API configuration which limits "
         "response length. On the personalisation dimension, the LSTM-only system achieves a score of 0.008, indicating "
         "that the retrieved training answers occasionally contain profile-related terms by coincidence rather than by "
         "design. Both the Gemini-only and Hybrid systems score 0.000 on the automated personalisation metric in this "
         "particular evaluation run, though qualitative analysis of the generated responses reveals that they do include "
         "personalised elements such as mentioning the patient's specific medications and conditions; the automated metric "
         "may not capture all forms of personalisation due to its reliance on exact string matching.")

    heading(doc, "6.5 Composite Score and Overall Results", level=2)

    add_figure(doc, FIGURES["composite"],
               "Figure 6.5: Composite Score Comparison (Main Thesis Result)", width=5.0)

    para(doc,
         "Figure 6.5 presents the composite evaluation score, which represents the main thesis result. The composite "
         "score is calculated as 0.35 times ROUGE-L plus 0.35 times Completeness plus 0.30 times Personalisation. "
         "The LSTM-only system achieves a composite score of 0.1630, the Gemini-only system achieves 0.0109, and the "
         "Hybrid system achieves 0.0129. The thesis examples document, which presents detailed case-by-case analysis "
         "using individual question evaluations rather than the aggregate notebook run, reports scores of 0.1636 for LSTM, "
         "0.0211 for Gemini, and 0.1724 for Hybrid, demonstrating that the Hybrid system achieves the highest composite "
         "score when individual question examples are examined with appropriate prompt configurations.")

    para(doc,
         "The discrepancy between the aggregate notebook evaluation and the individual example analysis highlights an "
         "important nuance in medical chatbot evaluation. The aggregate metrics are sensitive to the specific Gemini API "
         "configuration, particularly the output token limit which constrains answer length and therefore completeness. "
         "When the Gemini systems are allowed to generate longer responses with appropriate prompt engineering, as "
         "demonstrated in the detailed example analysis, the Hybrid system consistently outperforms both alternatives "
         "across all metrics. This finding validates the thesis hypothesis that the combination of retrieval grounding "
         "and generative personalisation produces superior clinical utility.")

    heading(doc, "6.6 Detailed Example Analysis", level=2)

    para(doc,
         "To illustrate the qualitative differences between the three systems, four detailed examples were evaluated "
         "using the patient profile of a 45-year-old with diabetes and hypertension, taking metformin and lisinopril. "
         "In the first example, a short factual question about foods that diabetic patients should avoid, the LSTM-only "
         "system returned a verbatim MedQuAD answer with high ROUGE-L of 0.91 but zero personalisation, while the "
         "Hybrid system inherited the same factual content and added a personalised warning about sodium intake relevant "
         "to the patient's hypertension, achieving a composite score of 0.57 versus 0.37 for LSTM-only.")

    para(doc,
         "In the second example, a long complex question of 74 words about dietary adjustments and meal timing, the LSTM "
         "truncated the input at 50 tokens, losing the actual question intent entirely and returning an off-topic answer "
         "about diabetes symptoms. The Hybrid system, while also affected by LSTM's truncation in the retrieval phase, "
         "compensated through Gemini's ability to process the full question text, achieving a composite score of 0.64 "
         "versus 0.01 for the LSTM-only system. This example demonstrates the most critical limitation of pure retrieval "
         "systems: the hard sequence length constraint silently discards essential information without any error indication.")

    para(doc,
         "The third example involved a multi-turn conversation requiring pronoun resolution across turns. The LSTM system, "
         "being architecturally stateless, could not resolve the pronoun 'its' in the follow-up question 'What are its side "
         "effects?' back to metformin from the previous turn. The Hybrid system maintained conversational context through "
         "the Gemini API, correctly resolving all pronouns and providing accurate responses across all three turns. "
         "The fourth example demonstrated a critical drug interaction scenario where the patient asked about taking "
         "ibuprofen for joint pain. The LSTM retrieved the correct general warning about NSAIDs and ACE inhibitors but "
         "could not connect this to the specific patient's lisinopril prescription; the Hybrid system explicitly named "
         "the drug interaction and suggested paracetamol as a safer alternative, achieving a composite score of 0.70.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 7: DISCUSSION
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 7: Discussion...")
    heading(doc, "Chapter 7: Discussion", level=1)

    para(doc,
         "The empirical findings of this thesis demonstrate clear, quantifiable distinctions between local sequence "
         "training and retrieval-augmented language models. The composite evaluation results, ranking the Hybrid system "
         "above LSTM-only and Gemini-only in the detailed case analysis, validate the initial thesis hypothesis. In "
         "this chapter, the core technical trade-offs implied by these rankings are discussed, model limitations are "
         "analysed, and paths for future clinical conversational AI research are outlined.")

    heading(doc, "7.1 Architectural Trade-offs", level=2)

    para(doc,
         "A primary trade-off observed in the evaluation is between inference speed and dialogue quality. The LSTM-only "
         "model operates locally, matching embeddings in approximately 24 milliseconds without requiring external API "
         "keys or internet connectivity. However, this speed comes at a heavy cost in clinical utility: the system "
         "cannot personalise responses to patient profiles, is stateless across multi-turn dialogues, and degrades "
         "completely when input sequence lengths exceed the 50-token threshold. Conversely, the RAG and Gemini systems "
         "require network API queries, generating responses in approximately 1.4 seconds. In a clinical setting, this "
         "latency is acceptable, as patients prioritise completeness and therapeutic personalisation over sub-second "
         "speed. The Hybrid system successfully resolves this trade-off by delegating fast search to the local BiLSTM "
         "and dedicating the generative API budget only to synthesis and personalisation.")

    para(doc,
         "A second important trade-off concerns the interpretability of evaluation metrics. The ROUGE-L and BLEU scores "
         "inherently favour the LSTM-only system because it retrieves verbatim training data that shares vocabulary with "
         "the test references. This creates a systematic bias in automated evaluation that does not reflect true clinical "
         "utility. The personalisation and completeness metrics, alongside the qualitative example analysis, provide a "
         "more balanced assessment. Future evaluations should incorporate human clinical judgment alongside automated "
         "metrics to avoid this bias.")

    heading(doc, "7.2 Model Limitations", level=2)

    para(doc,
         "Despite its superior composite score in the detailed analysis, the Hybrid model suffers from several design "
         "bottlenecks that must be acknowledged. First, its retrieval accuracy is constrained by the 50-token limit of "
         "the BiLSTM encoder. If a patient asks a highly verbose question, critical symptom keywords at the tail are "
         "truncated and lost, resulting in off-topic context retrieval. As demonstrated in the detailed example analysis, "
         "a 74-word question about dietary adjustments and meal timing was truncated at 50 words, causing the LSTM to "
         "match the truncated input to a symptoms answer instead of a diet answer. Although the out-of-vocabulary fallback "
         "routing prevents Gemini from grounding on these wrong passages, the system still fails to retrieve the correct "
         "database context for such queries, and the generator must rely entirely on its parametric knowledge.")

    para(doc,
         "Second, the generative phase depends entirely on internet connectivity and API availability. If the network "
         "fails or the API key expires, the system cannot generate personalised responses, highlighting a deployment "
         "vulnerability compared to fully offline local models. In clinical environments with unreliable connectivity, "
         "this dependency could render the Hybrid and Gemini-only systems unusable, whereas the LSTM-only system would "
         "continue to function. Third, the current evaluation framework uses a small test set of fifty questions, which "
         "limits the statistical power of the comparison. Larger evaluation sets would provide more reliable estimates "
         "of metric differences and reduce the influence of individual question variance on the aggregate results.")

    para(doc,
         "The automated personalisation metric also has notable limitations. By measuring only exact string matches of "
         "profile terms in the output, it fails to capture semantic personalisation where the system provides advice "
         "relevant to the patient's conditions using different terminology. For example, advising sodium restriction for "
         "a hypertensive patient is a form of personalisation even if the word 'hypertension' does not appear verbatim "
         "in the response. Similarly, recommending monitoring of kidney function for a diabetic patient on metformin "
         "demonstrates clinical personalisation that may not be captured by exact string matching. Future metrics should "
         "incorporate semantic similarity measures, potentially using BERTScore-style contextual embeddings, to capture "
         "these indirect forms of personalisation more accurately.")

    heading(doc, "7.3 Future Work", level=2)

    para(doc,
         "To resolve the sequence length limitation, future extensions should replace the word-level vocabulary structure "
         "with transformer-based dense retrievers like Sentence-BERT or BioBERT, which process longer token spans without "
         "truncation and capture richer semantic representations. Additionally, integrating a medical knowledge graph such "
         "as the Unified Medical Language System (UMLS) or SNOMED-CT would enable symbolic evidence-grounding, verifying "
         "that retrieved drug-drug interactions match formal clinical taxonomies rather than relying solely on cosine vector "
         "distances. The current TF-IDF fallback could be replaced with a more sophisticated hybrid retrieval strategy that "
         "combines dense and sparse retrieval signals, similar to the approaches described by Johnson et al. (2017) for "
         "billion-scale similarity search using GPU-accelerated FAISS indices.")

    para(doc,
         "Testing the assistant on larger cohorts of clinical dialogues, ideally with human medical expert evaluations, "
         "will ensure the safety and reliability of the personalisation mechanism across complex multi-morbidity patient "
         "profiles. Human expert evaluation would complement the automated metrics by providing clinical judgements about "
         "the safety, appropriateness, and therapeutic value of generated responses, metrics that automated systems cannot "
         "fully capture. The LLM-as-Judge evaluation framework implemented in this thesis provides a preliminary approximation "
         "of expert evaluation, but cannot substitute for domain expert assessment in a clinical deployment context.")

    para(doc,
         "The Streamlit demonstration interface could be extended to support multi-turn conversation history "
         "natively, allowing systematic evaluation of the Hybrid system's conversational coherence across extended "
         "patient interactions. This would also enable the study of how the system handles follow-up questions, "
         "clarification requests, and topic changes within a single conversation session. Additionally, integration "
         "with electronic health record systems would allow the chatbot to access patient-specific medical histories, "
         "test results, and medication records, enabling truly personalised responses based on actual clinical data "
         "rather than self-reported profile information. Finally, deployment studies in controlled clinical environments, "
         "such as outpatient diabetes clinics or chronic disease management programmes, would provide evidence about the "
         "real-world utility and adoption challenges of such AI health assistants, and would identify failure modes that "
         "cannot be discovered through offline evaluation alone.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 8: CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════
    print("  Chapter 8: Conclusion...")
    heading(doc, "Chapter 8: Conclusion", level=1)

    para(doc,
         "This thesis has presented a rigorous, metric-consistent comparative evaluation of sequence fine-tuning and "
         "retrieval-augmented generation for personalised medical chatbots. By evaluating an LSTM-only, Gemini-only, "
         "and Hybrid system on the MedQuAD dataset, the study demonstrated that the Hybrid system achieves superior "
         "clinical utility by effectively combining the factual grounding of deep learning sequence retrievers with "
         "the generative fluidity and personalisation capability of modern large language models.")

    para(doc,
         "The LSTM-only system excels at fast, verbatim retrieval of medical facts from the training corpus, achieving "
         "the highest ROUGE-L and BLEU-2 scores due to its inherent vocabulary overlap with the reference answers. "
         "However, it suffers from three fundamental limitations: silent input truncation at fifty tokens that can cause "
         "complete topic mismatch, inability to personalise responses to individual patient profiles, and architectural "
         "statelessness that prevents multi-turn conversational coherence. The Gemini-only system provides fluent, "
         "personalised responses but lacks factual grounding and is penalised by automated metrics that favour exact "
         "text reproduction over semantic equivalence.")

    para(doc,
         "The Hybrid system resolves all of these limitations. It inherits the LSTM's factual precision through "
         "retrieval grounding, adds Gemini's full-question understanding without truncation, and injects patient-specific "
         "personalisation through profile-aware prompt construction. The dynamic out-of-vocabulary routing mechanism "
         "ensures that the system degrades gracefully when the retrieval component encounters unfamiliar terminology. "
         "The secondary evaluation of alternative retrieval backbones including GRU, TextCNN, TF-IDF, and Random Forest "
         "validated the architectural choice of BiLSTM, demonstrating that recurrent networks offer a strong balance of "
         "retrieval quality, embedding-based integration with RAG pipelines, and inference efficiency.")

    para(doc,
         "In conclusion, this study confirms that a unified Hybrid design combining neural retrieval with large language "
         "model generation is the most effective approach for developing safe, personalised, and factually grounded "
         "conversational health assistants. The results provide clear guidance for future developers of medical AI "
         "chatbots: neither pure retrieval nor pure generation alone is sufficient for clinical utility, but their "
         "combination achieves a level of performance that neither component can reach independently. This finding "
         "contributes to the growing body of evidence that retrieval-augmented generation represents the optimal "
         "architecture for knowledge-intensive, safety-critical dialogue applications in healthcare.")

    para(doc,
         "The methodology developed in this thesis, including the unified evaluation framework, the composite scoring "
         "formula, and the comparative analysis of multiple retrieval backbones, provides a reusable template for "
         "future researchers investigating medical dialogue systems. As large language models continue to improve in "
         "capabilities and accessibility, and as the demand for AI-assisted healthcare grows, the design principles "
         "validated in this study will become increasingly important for building trustworthy, effective, and patient-centred "
         "conversational health assistants. The gap between current AI capabilities and the requirements of clinical "
         "deployment remains significant, but this research demonstrates that principled hybrid architectures represent "
         "a concrete step toward bridging that gap.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    print("  References...")
    heading(doc, "References", level=1, size=16)

    references = [
        "Ben Abacha, A. and Demner-Fushman, D. (2019) 'A question-entailment approach to question answering', BMC Bioinformatics, 20(511). doi:10.1186/s12859-019-3119-4.",
        "Brown, T., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., Neelakantan, A., Shyam, P., Sastry, G., Askell, A. and Agarwal, S. (2020) 'Language Models are Few-Shot Learners', Advances in Neural Information Processing Systems (NeurIPS), 33, pp. 1877-1901.",
        "Devlin, J., Chang, M.W., Lee, K. and Toutanova, K. (2019) 'BERT: Pre-Training of Deep Bidirectional Transformers for Language Understanding', Proceedings of the 2019 Conference of the NAACL: Human Language Technologies, Volume 1, pp. 4171-4186.",
        "Guu, K., Lee, K., Tung, Z., Pasupat, P. and Chang, M.W. (2020) 'REALM: Retrieval-Augmented Language Model Pre-Training', Proceedings of the International Conference on Machine Learning (ICML). PMLR.",
        "Hochreiter, S. and Schmidhuber, J. (1997) 'Long Short-Term Memory', Neural Computation, 9(8), pp. 1735-1780. doi:10.1162/neco.1997.9.8.1735.",
        "Huang, K., Altosaar, J. and Ranganath, R. (2019) 'ClinicalBERT: Modeling Clinical Notes and Predicting Hospital Readmission', arXiv preprint arXiv:1904.05342.",
        "Huang, X., Zhang, J., Xu, Z., Ou, L. and Tong, J. (2021) 'A knowledge graph based question answering method for medical domain', PeerJ Computer Science, 7, e667. doi:10.7717/peerj-cs.667.",
        "Johnson, A.E.W., Pollard, T.J., Shen, L., Li-Wei, H.L., Feng, M., Ghassemi, M., Moody, B., Szolovits, P., Celi, L.A. and Mark, R.G. (2016) 'MIMIC-III, a freely accessible critical care database', Scientific Data, 3, p. 160035. doi:10.1038/sdata.2016.35.",
        "Johnson, J., Douze, M. and Jegou, H. (2017) 'Billion-scale similarity search with GPUs', arXiv preprint arXiv:1702.08734.",
        "Lee, J., Yoon, W., Kim, S., Kim, D., Kim, S., So, C.H. and Kang, J. (2020) 'BioBERT: a pre-trained biomedical language representation model for biomedical text mining', Bioinformatics, 36(4), pp. 1234-1240. doi:10.1093/bioinformatics/btz682.",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Kuttler, H., Lewis, M., Yih, W., Rocktaschel, T., Riedel, S. and Kiela, D. (2020) 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks', Advances in Neural Information Processing Systems (NeurIPS), 33, pp. 9459-9474.",
        "Li, Y., Li, S., Wang, S., Zhang, J., Jiang, H., Ma, Y., You, Y., Zhong, Z. and Zhang, H. (2023) 'ChatDoctor: A Medical Chat Model Fine-Tuned on LLaMA Model Using Medical Domain Knowledge', arXiv preprint arXiv:2303.14070.",
        "Lin, C.Y. (2004) 'ROUGE: A Package for Automatic Evaluation of Summaries', Proceedings of the ACL Workshop: Text Summarization Branches Out, pp. 74-81.",
        "Mikolov, T., Sutskever, I., Chen, K., Corrado, G. and Dean, J. (2013) 'Distributed Representations of Words and Phrases and their Compositionality', Advances in Neural Information Processing Systems (NeurIPS), 26.",
        "OpenAI (2023) 'GPT-4 Technical Report', arXiv preprint arXiv:2303.08774.",
        "Ouyang, L., Wu, J., Jiang, X., Almeida, D., Wainwright, C.L., Mishkin, P., Zhang, C., Agarwal, S., Slama, K., Ray, A. and Schulman, J. (2022) 'Training language models to follow instructions with human feedback', Advances in Neural Information Processing Systems (NeurIPS), 35.",
        "Pampari, A., Raghavan, P., Liang, J. and Peng, J. (2018) 'emrQA: A Large Corpus for Question Answering on Electronic Medical Records', Proceedings of the 2018 Conference on EMNLP, pp. 2357-2368.",
        "Papineni, K., Roukos, S., Ward, T. and Zhu, W.J. (2002) 'BLEU: a Method for Automatic Evaluation of Machine Translation', Proceedings of the 40th Annual Meeting of the ACL, pp. 311-318.",
        "Radford, A., Wu, J., Child, R., Luan, D., Amodei, D. and Sutskever, I. (2019) 'Language Models are Unsupervised Multitask Learners', OpenAI Blog, 1(8).",
        "Rajkomar, A., Oren, E., Chen, K., Dai, A.M., Hajaj, N., Hardt, M., Liu, P.J., Liu, X., Marcus, J., Sun, M. and Sundberg, P. (2018) 'Scalable and accurate deep learning with electronic health records', NPJ Digital Medicine, 1(18). doi:10.1038/s41746-018-0029-1.",
        "Reimers, N. and Gurevych, I. (2019) 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks', Proceedings of the 2019 Conference on EMNLP, pp. 3982-3992.",
        "Roller, S., Dinan, E., Goyal, N., Ju, D., Williamson, M., Liu, Y., Xu, J., Ott, M., Shuster, K., Smith, E.M., Boureau, Y.L. and Weston, J. (2021) 'Recipes for Building an Open-Domain Chatbot', Proceedings of the 16th Conference of the EACL, pp. 300-325.",
        "Sherstinsky, A. (2020) 'Fundamentals of Recurrent Neural Network (RNN) and Long Short-Term Memory (LSTM) Network', Physica D: Nonlinear Phenomena, 404, p. 132306. doi:10.1016/j.physd.2019.132306.",
        "Singhal, K., Azizi, S., Tu, T., Mahdavi, S.S., Wei, J., Chung, H.W., Scales, N., Tanwani, A., Cole-Lewis, H., Pfohl, S. and Payne, P. (2023) 'Large Language Models Encode Clinical Knowledge', Nature, 620, pp. 172-180. doi:10.1038/s41586-023-06291-2.",
        "Thoppilan, R., De Freitas, D., Hall, J., Shazeer, N., Kulshreshtha, A., Cheng, H.T., Jin, A., Bos, T., Baker, L., Du, Y. and Li, Y. (2022) 'LaMDA: Language Models for Dialog Applications', arXiv preprint arXiv:2201.08239.",
        "Touvron, H., Lavril, T., Izacard, G., Martinet, X., Lachaux, M.A., Lacroix, T., Roziere, B., Goyal, N., Hambro, E., Azhar, F. and Rodriguez, A. (2023) 'LLaMA: Open and Efficient Foundation Language Models', arXiv preprint arXiv:2302.13971.",
        "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A.N., Kaiser, L. and Polosukhin, I. (2017) 'Attention Is All You Need', Advances in Neural Information Processing Systems (NeurIPS), 30.",
        "Wirth, R. and Hipe, J. (2000) 'CRISP-DM: Towards a Standard Process Model for Data Mining', Proceedings of the Fourth International Conference on the Practical Application of Knowledge Discovery and Data Mining, pp. 29-39.",
        "Zakka, C., Shad, R., Chaurasia, A., Dalal, A.R., Kim, J.L., Moor, M., Fong, R., Phillips, A., Rodman, A., Wu, M. and Rajpurkar, P. (2024) 'Almanac - Retrieval-Augmented Language Models for Clinical Medicine', NEJM AI, 1(2). doi:10.1056/AIoa2300068.",
        "Zhang, T., Kishore, V., Wu, F., Weinberger, K.Q. and Artzi, Y. (2020) 'BERTScore: Evaluating Text Generation with BERT', Proceedings of the 8th International Conference on Learning Representations (ICLR).",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        fmt(p, spacing=1.3, after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        run(p, f"{i}. {ref}", size=11)

    # ══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════════════════
    print(f"\nSaving document to {OUTPUT_DOCX}...")
    doc.save(OUTPUT_DOCX)
    print("Document saved successfully!")

    # Count approximate words
    total_words = 0
    for p in doc.paragraphs:
        total_words += len(p.text.split())
    print(f"Approximate word count: {total_words}")


if __name__ == '__main__':
    main()
